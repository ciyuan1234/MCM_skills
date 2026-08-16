#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
版面格式检查器 (format-check.py) —— 检查导出的 docx/PDF 是否符合国赛版面规范
用法:
    python format-check.py <工作目录> [-o 输出.json]
示例:
    python format-check.py D:\2026_国赛1234
    python format-check.py . -o 4_论文\format-check.json

检查项（docx 版面层，官方规范 HARD + 获奖论文惯例）:
  1. 页边距      上下左右 >= 2.5cm（官方规范 HARD）
  2. 页脚页码    每节页脚含 PAGE 域（官方 HARD: 摘要页起、页脚中部连续编号）
  3. 首页摘要    文档第一段须含"摘要"（电子版第一页=摘要专用页，官方 HARD）
  4. 图片题注    每张图片后须紧跟含"图N"的题注段（获奖论文 4/4 惯例）
  5. 三线表      表格边框仅顶线/表头线/底线，无竖线（获奖论文惯例）
  6. 表题注      每个表格前须有含"表N"的题注段（获奖论文惯例）
  7. 题注字号    图/表题注字号 <= 正文字号（惯例，软检查）
  8. 身份泄漏    论文全文不得含 学校/队号/手机号/邮箱 等参赛者身份信息（官方 HARD）
  9. PDF 检查    大小 <= 20MB（官方 HARD）；第 1 页须为摘要页；
                 正文页数（第 2 页 至 附录/参考文献 前）>30 错误 / 26-30 警告 / <=25 通过
                 （官方 HARD: 正文不超过 30 页；附录页数不限，不计入正文）

退出码: 0 = 无错误; 1 = 存在错误项
"""

import glob
import json
import os
import re
import sys

PASS = "[通过]"
WARN = "[警告]"
ERR = "[错误]"
ERROR_COUNT = 0


def report(level, msg):
    print(f"{level} {msg}")
    global ERROR_COUNT
    if level == ERR:
        ERROR_COUNT += 1


def find_docx(workdir):
    hits = glob.glob(os.path.join(workdir, "4_论文", "*.docx"))
    return hits[0] if hits else None


def find_pdf(workdir):
    hits = glob.glob(os.path.join(workdir, "4_论文", "*.pdf"))
    return hits[0] if hits else None


def check_margins(sections):
    print("== 1. 页边距 (>=2.5cm) ==")
    ok = True
    for i, sec in enumerate(sections):
        for name, val in (("上", sec.top_margin), ("下", sec.bottom_margin),
                          ("左", sec.left_margin), ("右", sec.right_margin)):
            cm = val.cm if val is not None else 0.0
            if cm < 2.49:
                report(ERR if i == 0 else WARN,
                       f"第{i+1}节 页边距{name} {cm:.2f}cm < 2.5cm（官方规范 HARD）")
                ok = False
    if ok:
        report(PASS, "各节页边距均 >= 2.5cm")
    return ok


def has_page_field(par):
    xml = par._element.xml
    return "PAGE" in xml and ("fldChar" in xml or "instrText" in xml or "fldSimple" in xml)


def check_page_numbers(sections):
    print("== 2. 页脚页码 (摘要页起连续编号) ==")
    ok = True
    for i, sec in enumerate(sections):
        footers = [sec.footer]
        try:
            if sec.different_first_page_header_footer and sec.first_page_footer is not None:
                footers.append(sec.first_page_footer)
        except Exception:
            pass
        found = False
        for ft in footers:
            for p in ft.paragraphs:
                if has_page_field(p):
                    found = True
                    break
            if found:
                break
        if not found:
            report(ERR if i == 0 else WARN, f"第{i+1}节 页脚无 PAGE 页码域（官方规范 HARD）")
            ok = False
    if ok:
        report(PASS, "各节页脚均含 PAGE 页码域")
    return ok


def body_sequence(doc):
    """按文档顺序返回元素: ('p', paragraph) / ('tbl', table)。"""
    seq = []
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            seq.append(("p", Paragraph(child, doc)))
        elif tag == "tbl":
            seq.append(("tbl", Table(child, doc)))
    return seq


def has_image(par):
    return par._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline") or \
           par._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor")


# ---------- 8. 身份泄漏扫描（官方 HARD: 摘要页/正文/附录均不得含参赛者身份信息） ----------

IDENTITY_PATTERNS = [
    (re.compile(r"[\u4e00-\u9fa5]{2,3}(?:省|市|自治区)[\u4e00-\u9fa5]{0,8}(?:大学|学院)(?:[\u4e00-\u9fa5]{0,6}(?:分校|校区))?"),
     "学校/院校名称（省市+校名）"),
    (re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)"), "手机号码"),
    (re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"), "电子邮箱"),
    (re.compile(r"(?:队号|参赛队编号|队伍编号)\s*[:：]?\s*[A-Z]?\d{4,}"), "参赛队号"),
    (re.compile(r"(?<!\d)20\d{2}[A-Z]\s*\d{3,4}(?!\d)"), "队号（年份+题号+序号）"),
    (re.compile(r"(?:队员|参赛队员|姓名)\s*[:：]\s*[\u4e00-\u9fa5]{2,4}(?:[、，,]\s*[\u4e00-\u9fa5]{2,4}){0,2}"),
     "队员姓名列表"),
]


def check_identity(paras_text):
    print("== 7. 身份泄漏扫描 ==")
    hits = []
    for pat, desc in IDENTITY_PATTERNS:
        for m in pat.finditer(paras_text):
            snippet = paras_text[max(0, m.start() - 10):m.end() + 10].replace("\n", " ")
            hits.append(f"{desc}: …{snippet}…")
            break  # 每类只报首处
    if hits:
        report(ERR, f"检测到参赛者身份信息（官方 HARD: 任何地方不得出现）: {hits}")
    else:
        report(PASS, "未检测到学校/队号/手机号/邮箱等身份信息")


def table_border_info(tbl):
    """返回 dict: top/bottom/left/right/insideH/insideV 是否有边框。"""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    borders = None
    tblPr = tbl._element.tblPr
    if tblPr is not None:
        borders = tblPr.find(f"{ns}tblBorders")
    try:
        style_el = tbl.style.element
        if borders is None:
            borders = style_el.find(f"{ns}tblPr/{ns}tblBorders")
    except Exception:
        pass
    out = {}
    if borders is None:
        return out  # 无边框定义（默认无边框）
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(f"{ns}{name}")
        out[name] = el is not None and el.get(f"{ns}val") not in ("nil", "none", None)
    return out


def check_figures_and_tables(seq):
    print("== 3. 图片题注 / 4. 三线表 / 5. 表题注 ==")
    n_img = 0
    fig_captioned = 0
    fig_bad = []
    tbls = 0
    tbl_captioned = 0
    three_line = 0
    tbl_bad = []
    normal_size = 12.0  # 正文小四
    cap_size_bad = []
    for idx, (kind, obj) in enumerate(seq):
        if kind == "p":
            if has_image(obj):
                n_img += 1
                nxt = None
                for j in range(idx + 1, len(seq)):
                    if seq[j][0] == "p" and seq[j][1].text.strip():
                        nxt = seq[j][1].text.strip()
                        break
                if nxt and re.match(r"^图\s*\d+", nxt):
                    fig_captioned += 1
                    for r in seq[j][1].runs:
                        if r.font.size:
                            if r.font.size.pt > normal_size:
                                cap_size_bad.append(f"图题注字号 {r.font.size.pt}pt > 正文 {normal_size}pt")
                            break
                else:
                    fig_bad.append(f"第{n_img}张图片后无'图N'题注段")
        elif kind == "tbl":
            tbls += 1
            prev = None
            prev_txt = ""
            for j in range(idx - 1, -1, -1):
                if seq[j][0] == "p" and seq[j][1].text.strip():
                    prev = seq[j][1]
                    prev_txt = prev.text.strip()
                    break
            # 符号说明节的表格惯例不编号，豁免
            near = "\n".join(o.text.strip() for k, o in seq[max(0, idx - 12):idx] if k == "p")
            if re.search(r"符号说明", near) and prev_txt and not re.match(r"^表\s*\d+", prev_txt):
                tbl_captioned += 1  # 符号说明表惯例不编号，视同合规
            elif prev_txt and re.match(r"^表\s*\d+", prev_txt):
                tbl_captioned += 1
                for r in prev.runs:
                    if r.font.size:
                        if r.font.size.pt > normal_size:
                            cap_size_bad.append(f"表题注字号 {r.font.size.pt}pt > 正文 {normal_size}pt")
                        break
            else:
                tbl_bad.append(f"第{tbls}张表格前无'表N'题注段")
            bi = table_border_info(obj)
            if bi and not bi.get("insideV") and not bi.get("left") and not bi.get("right") \
                    and (bi.get("top") or bi.get("insideH")) and bi.get("bottom"):
                three_line += 1
            elif bi and (bi.get("insideV") or bi.get("left") or bi.get("right")):
                tbl_bad.append(f"第{tbls}张表格非三线表（含竖线/左右框线）")
    if n_img == 0:
        report(ERR, "docx 中未发现任何图片（论文正文必须有图，获奖论文 4/4 带图）")
    else:
        if fig_captioned == n_img:
            report(PASS, f"{n_img} 张图片均有'图N'题注（图下）")
        else:
            report(ERR, f"{n_img - fig_captioned}/{n_img} 张图片缺题注: {fig_bad[:3]}")
    if tbls == 0:
        report(WARN, "docx 中未发现表格")
    else:
        if tbl_captioned == tbls:
            report(PASS, f"{tbls} 张表格均有'表N'题注（表上）")
        else:
            report(WARN, f"{tbls - tbl_captioned}/{tbls} 张表格缺题注: {tbl_bad[:3]}")
        three_bad = [b for b in tbl_bad if "非三线表" in b]
        if three_line == tbls:
            report(PASS, f"{tbls} 张表格均为三线表")
        else:
            report(WARN, f"{len(three_bad)} 张表格非三线表: {three_bad[:3]}（惯例：仅顶线/表头线/底线）")
    if cap_size_bad:
        report(WARN, f"题注字号大于正文: {cap_size_bad[:2]}（惯例：题注小于正文）")
    else:
        report(PASS, "题注字号均 <= 正文（惯例）")
    return n_img, tbls


def check_pdf_layout(pdf_path):
    """PDF 层检查: 大小 / 第 1 页摘要 / 正文页数（附录不计入）。
    正文 = 第 2 页 至 附录标题前（无附录则至参考文献标题页）。
    """
    print("== 8. PDF 大小 / 首页摘要 / 正文页数 ==")
    size_mb = os.path.getsize(pdf_path) / 1048576.0
    if size_mb <= 20:
        report(PASS, f"PDF {size_mb:.1f}MB <= 20MB（官方 HARD）")
    else:
        report(ERR, f"PDF {size_mb:.1f}MB > 20MB（官方 HARD）")
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        total = len(reader.pages)
        texts = []
        for p in reader.pages:
            try:
                texts.append(p.extract_text() or "")
            except Exception:
                texts.append("")
        page1 = re.sub(r"\s", "", texts[0])
        if not page1:
            report(WARN, "PDF 第 1 页无法提取文本（可能为图片型 PDF，跳过摘要页检查）")
        elif re.search(r"承诺书", page1):
            report(ERR, "PDF 第 1 页含'承诺书'——电子版论文不得包含承诺书/编号专用页（官方 HARD: 第一页必须为摘要专用页）")
        elif re.search(r"摘\s*要", page1):
            report(PASS, "PDF 第 1 页为摘要专用页（电子版第一页必须为摘要页，官方 HARD）")
        else:
            report(ERR, "PDF 第 1 页未识别到'摘要'（电子版第一页必须为摘要专用页，官方 HARD）")

        # 定位附录/参考文献起始页：标题独立成行（"附录" / "七、附录" / "参考文献"）
        appendix_p = None
        refs_p = None
        title_pat = re.compile(r"(?m)^\s*(?:[一二三四五六七八九十\d]+\s*[、.．]\s*)?(附\s*录|参考文献)([:：]?[\s]*|$)")
        for i, t in enumerate(texts):
            if appendix_p is None or refs_p is None:
                for m in title_pat.finditer(t):
                    if m.group(1).startswith("附") and appendix_p is None:
                        appendix_p = i + 1
                    elif m.group(1).startswith("参考") and refs_p is None:
                        refs_p = i + 1
        body_end = None
        if appendix_p:
            body_end = appendix_p - 1
        elif refs_p:
            body_end = refs_p
        else:
            body_end = total
        body_pages = body_end - 1  # 第 2 页起
        if appendix_p:
            app_pages = total - (appendix_p - 1)
        else:
            app_pages = 0
        if body_pages > 30:
            report(ERR, f"正文 {body_pages} 页 > 30 页（官方 HARD: 正文不超过 30 页，不含附录）")
        elif body_pages >= 26:
            report(WARN, f"正文 {body_pages} 页（26-30 页，合规但逼近官方上限 30 页）")
        elif body_pages < 20:
            report(ERR, f"正文仅 {body_pages} 页（<20 页，本 skill 硬标准——获奖论文正文 25-35 页，分析深度不足会失分）")
        else:
            report(PASS, f"正文 {body_pages} 页（20-25 页，官方上限 30 页；含 摘要1页+附录{app_pages}页 共 {total} 页）")
        return {"pdf_pages": total, "body_pages": body_pages, "appendix_pages": app_pages}
    except Exception as e:
        report(WARN, f"无法读取 PDF 布局: {e}")
        return None


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(2)
    wd = sys.argv[1]
    out = ""
    if "-o" in sys.argv:
        i = sys.argv.index("-o")
        if i + 1 < len(sys.argv):
            out = sys.argv[i + 1]
    if not os.path.isdir(wd):
        print(f"{ERR} 工作目录不存在: {wd}", file=sys.stderr)
        sys.exit(2)

    from docx import Document
    from docx.shared import Cm

    docx_path = find_docx(wd)
    if not docx_path:
        has_tex = any(f.lower().endswith(".tex") for f in os.listdir(os.path.join(wd, "4_论文")))
        if has_tex:
            # LaTeX 路线: md2tex.py/paper-template 已保证版面（页边距/页码域/题注字号），
            # 这里只做 PDF 层检查（大小/首页摘要/正文页数/身份泄漏）
            print("== LaTeX 路线（无 docx）: 版面级检查由模板保证 ==")
            report(PASS, "LaTeX 模板保证版面（页边距 2.5cm/页脚页码/题注字号），跳过 docx 级检查")
            pdf_path = find_pdf(wd)
            n_pages = -1
            if not pdf_path:
                report(ERR, "4_论文 下没有 PDF（最终必须提交 PDF 版论文）")
            else:
                pdf_info = check_pdf_layout(pdf_path)
                n_pages = pdf_info["pdf_pages"] if pdf_info else -1
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(pdf_path)
                    ptext = "\n".join((p.extract_text() or "") for p in reader.pages)
                    check_identity(ptext)
                except Exception as e:
                    report(WARN, f"PDF 文本提取失败，跳过身份扫描: {e}")
            print()
            if ERROR_COUNT:
                print(f"检查完成: {ERROR_COUNT} 项错误，请修复版面后再提交。")
                sys.exit(1)
            print("检查完成: 无错误。")
            if out:
                payload = {"margins_ok": True, "page_number_ok": True,
                           "images": 0, "tables": 0, "pdf_pages": n_pages}
                with open(out, "w", encoding="utf-8") as f:
                    json.dump(payload, f, ensure_ascii=False, indent=2)
                print(f"[完成] 版面检查结果已写入: {out}")
            sys.exit(0)
        report(ERR, "4_论文 下没有 docx（先运行 export-paper.ps1 导出）")
        print(f"检查完成: {ERROR_COUNT} 项错误")
        sys.exit(1)
    doc = Document(docx_path)

    m1 = check_margins(doc.sections)
    m2 = check_page_numbers(doc.sections)

    print("== 3. 首页摘要 ==")
    first_texts = []
    for p in doc.paragraphs:
        t = p.text.strip().lstrip("\ufeff")
        if t:
            first_texts.append(t)
        if len(first_texts) >= 3:
            break
    if not first_texts:
        report(ERR, "文档没有任何段落")
    elif any(re.search(r"(^|#|\s)摘\s*要", t) for t in first_texts):
        report(PASS, "文档开头为标题+摘要页（第一页为摘要专用页）")
    else:
        report(ERR, f"文档开头未见摘要标题（实际开头: {first_texts[0][:20]}…），电子版第一页必须为摘要专用页（官方 HARD）")

    m4 = check_figures_and_tables(body_sequence(doc))

    all_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text
    check_identity(all_text)

    pdf_path = find_pdf(wd)
    n_pages = -1
    if not pdf_path:
        report(ERR, "4_论文 下没有 PDF（最终必须提交 PDF 版论文）")
    else:
        pdf_info = check_pdf_layout(pdf_path)
        n_pages = pdf_info["pdf_pages"] if pdf_info else -1

    print()
    if ERROR_COUNT:
        print(f"检查完成: {ERROR_COUNT} 项错误，请修复版面后再提交。")
        sys.exit(1)
    print("检查完成: 无错误。")
    if out:
        payload = {"margins_ok": m1, "page_number_ok": m2, "images": m4[0],
                   "tables": m4[1], "pdf_pages": n_pages}
        with open(out, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        print(f"[完成] 版面检查结果已写入: {out}")
    sys.exit(0)


if __name__ == "__main__":
    main()
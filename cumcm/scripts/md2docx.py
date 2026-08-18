#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown -> docx 转换器 (md2docx.py)  [依赖 python-docx: pip install python-docx]
用法:
    python md2docx.py paper.md -o paper.docx

支持: 标题(#/##/###) / 段落 / 加粗 / 斜体 / 行内代码 / 表格 / 无序有序列表 /
      图片 ![alt](path) / 围栏代码块 / 引用。
字体: 正文宋体小四, 标题黑体（符合国赛排版惯例）。
"""

import argparse
import os
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[错误] 缺少 python-docx，请先执行: pip install python-docx", file=__import__("sys").stderr)
    raise SystemExit(2)


def setup_page(doc):
    """页边距 >= 2.5cm + 页脚居中 PAGE 页码（官方规范 HARD）。"""
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.size = Pt(10.5)
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        run._r.append(f1); run._r.append(it); run._r.append(f2)


def make_three_line_table(table):
    """三线表: 顶线/底线 1.5pt, 表头下 0.75pt, 无竖线（获奖论文惯例）。"""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for name, sz in (("top", "12"), ("bottom", "12")):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), sz); el.set(qn("w:color"), "000000")
        borders.append(el)
    for name in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}"); el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:color"), "000000")
        tcB.append(b)
        tcPr.append(tcB)


def caption_style(par, size_pt=11):
    """题注样式: 居中、宋体、小一号、加粗（惯例: 题注字号 <= 正文）。"""
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in par.runs:
        set_cn_font(run, "宋体", size_pt=size_pt, bold=True)


def set_cn_font(run, name_east, name_ascii="Times New Roman", size_pt=12, bold=False):
    run.font.name = name_ascii
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), name_east)


def add_runs_with_inline(par, text, base_bold=False):
    """解析 **bold** *italic* `code` 行内标记。"""
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|`[^`]*`)", text)
    for t in tokens:
        if not t:
            continue
        if t.startswith("**") and t.endswith("**") and len(t) > 4:
            run = par.add_run(t[2:-2]); set_cn_font(run, "宋体", bold=True or base_bold)
        elif t.startswith("*") and t.endswith("*") and len(t) > 2:
            run = par.add_run(t[1:-1]); run.font.italic = True; set_cn_font(run, "宋体", bold=base_bold)
        elif t.startswith("`") and t.endswith("`") and len(t) > 2:
            run = par.add_run(t[1:-1]); run.font.name = "Consolas"; run.font.size = Pt(11)
        else:
            run = par.add_run(t); set_cn_font(run, "宋体", bold=base_bold)


def convert(md_path, out_path):
    with open(md_path, encoding="utf-8-sig") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(7.5)
    style.paragraph_format.space_before = Pt(0)
    setup_page(doc)

    i = 0
    in_code = False
    code_buf = []
    last_caption_line = None  # 最近一个"表N"题注行（用于表格渲染后回设样式）
    while i < len(lines):
        ln = lines[i].rstrip()

        # 围栏代码块
        if ln.strip().startswith("```"):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False
                for cl in code_buf:
                    p = doc.add_paragraph()
                    run = p.add_run(cl)
                    run.font.name = "Consolas"; run.font.size = Pt(10.5)
                p = doc.paragraphs[-1]
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        if in_code:
            code_buf.append(ln)
            i += 1
            continue

        s = ln.strip()
        # 表格
        if s.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            header = [c.strip() for c in s.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            ncol = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=ncol)
            make_three_line_table(table)
            for c, txt in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].text = ""
                add_runs_with_inline(cell.paragraphs[0], txt, base_bold=True)
            for r, row in enumerate(rows, start=1):
                for c in range(ncol):
                    txt = row[c] if c < len(row) else ""
                    cell = table.rows[r].cells[c]
                    cell.paragraphs[0].text = ""
                    add_runs_with_inline(cell.paragraphs[0], txt)
            if last_caption_line and re.match(r"^表\s*\d+", last_caption_line):
                for p in reversed(doc.paragraphs):
                    if p.text.strip():
                        caption_style(p)
                        break
            doc.add_paragraph()
            i += 1
            continue

        # 标题
        hm = re.match(r"^(#{1,6})\s+(.*)$", s)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            p = doc.add_heading("", level=min(level, 4))
            run = p.add_run(text)
            set_cn_font(run, "黑体", size_pt=16 - level * 2, bold=True)
            i += 1
            continue

        # 图片
        im = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", s)
        if im:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            img_path = im.group(2)
            if not os.path.isabs(img_path):
                bases = [os.path.dirname(os.path.abspath(md_path)),
                         os.path.dirname(os.path.dirname(os.path.abspath(md_path))),
                         os.getcwd()]
                for b in bases:
                    cand = os.path.join(b, img_path)
                    if os.path.isfile(cand):
                        img_path = cand
                        break
            try:
                run.add_picture(img_path, width=Inches(5.5))
            except Exception as e:
                run2 = p.add_run(f"[图片缺失: {im.group(2)}]")
                set_cn_font(run2, "宋体")
            if im.group(1):
                cap = doc.add_paragraph()
                caption_style(cap, size_pt=11)
                cap_run = cap.add_run(im.group(1))
                set_cn_font(cap_run, "宋体", size_pt=11, bold=True)
            i += 1
            continue

        # 列表
        lm = re.match(r"^(\s*)[-*+]\s+(.*)$", s)
        if lm:
            p = doc.add_paragraph(style="List Bullet")
            if lm.group(1):
                p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (len(lm.group(1)) // 2))
            add_runs_with_inline(p, lm.group(2))
            i += 1
            continue
        om = re.match(r"^(\s*)\d+[.)]\s+(.*)$", s)
        if om:
            p = doc.add_paragraph(style="List Number")
            if om.group(1):
                p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (len(om.group(1)) // 2))
            add_runs_with_inline(p, om.group(2))
            i += 1
            continue

        # 引用
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            add_runs_with_inline(p, s.lstrip(">").strip())
            i += 1
            continue

        # 分隔线
        if re.match(r"^-{3,}$|^\*{3,}$", s):
            i += 1
            continue

        # 普通段落
        if s:
            p = doc.add_paragraph()
            add_runs_with_inline(p, s)
            if re.match(r"^表\s*\d+", s):
                last_caption_line = s
        i += 1

    doc.save(out_path)
    print(f"[完成] {md_path} -> {out_path}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("md")
    ap.add_argument("-o", "--out", default="")
    args = ap.parse_args()
    out = args.out or args.md.rsplit(".", 1)[0] + ".docx"
    convert(args.md, out)